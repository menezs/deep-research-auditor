from __future__ import annotations

import re
from pathlib import Path
from typing import Protocol

from ..common.errors import ExtractionError

_HEADING_STYLE_RE = re.compile(r"^heading\s*(\d)$", re.IGNORECASE)


def _heading_level(style_name: str) -> int | None:
    match = _HEADING_STYLE_RE.match(style_name.strip())
    if match:
        return min(int(match.group(1)), 6)
    return None


def _docx_run_markdown(run) -> str:
    """Converte um run do python-docx para markdown, preservando o unico
    sinal textual que distingue um marcador de citacao (`font.superscript`)
    de um numero comum de prosa ja achatada — sem isso, "bilheteira1" e
    indistinguivel de qualquer outro digito colado ao texto."""
    text = run.text
    if text and run.font.superscript:
        digits = re.findall(r"\d+", text)
        if digits:
            return "".join(f"[{d}]" for d in digits)
    return text


def _docx_paragraph_markdown(paragraph) -> str:
    """Reconstroi um paragrafo preservando marcadores de citacao
    (superscript -> `[N]`) e cabecalhos (estilo `Heading N` -> `#`*N).

    Usa `paragraph.iter_inner_content()` (nao `paragraph.runs`) porque
    `.runs` pula silenciosamente qualquer run dentro de `<w:hyperlink>` —
    confirmado que isso descartaria a URL de toda entrada da lista de
    fontes, que o Word/Google Docs grava como hyperlink, nao texto
    solto. `iter_inner_content()` devolve `Run`/`Hyperlink` em ordem, cada
    um com o texto correto."""
    parts: list[str] = []
    for item in paragraph.iter_inner_content():
        if hasattr(item, "runs"):  # Hyperlink: agrega seus proprios runs
            parts.append("".join(_docx_run_markdown(r) for r in item.runs) or item.text)
        else:
            parts.append(_docx_run_markdown(item))
    text = "".join(parts)

    level = _heading_level(paragraph.style.name) if paragraph.style is not None else None
    return f"{'#' * level} {text.strip()}" if level else text


def _docx_table_markdown(table) -> str:
    """Reconstroi uma tabela como blocos de texto (uma linha por bloco),
    reaproveitando `_docx_paragraph_markdown` por celula para preservar
    marcadores de citacao (superscript -> `[N]`) — sem isso, tabelas de
    resposta (comuns em respostas de Deep Research, ex: comparativos de
    bilheteira) desapareciam por completo do texto extraido, pois
    `document.paragraphs` nao inclui paragrafos dentro de tabelas."""
    rows_text: list[str] = []
    for row in table.rows:
        seen_cells: set[int] = set()
        cells_text: list[str] = []
        for cell in row.cells:
            if id(cell._tc) in seen_cells:  # celula duplicada por merge horizontal
                continue
            seen_cells.add(id(cell._tc))
            cell_text = "\n\n".join(_docx_paragraph_markdown(p) for p in cell.paragraphs if p.text.strip())
            if cell_text.strip():
                cells_text.append(cell_text)
        row_text = " ".join(cells_text)
        if row_text.strip():
            rows_text.append(row_text)
    return "\n\n".join(rows_text)


class AnswerLoader(Protocol):
    def load(self, path: Path) -> str: ...


class MarkdownAnswerLoader:
    def load(self, path: Path) -> str:
        return path.read_text(encoding="utf-8")


class PdfAnswerLoader:
    def load(self, path: Path) -> str:
        import pymupdf4llm

        return pymupdf4llm.to_markdown(str(path))


class DocxAnswerLoader:
    def load(self, path: Path) -> str:
        import docx
        from docx.table import Table

        document = docx.Document(str(path))
        blocks = [
            _docx_table_markdown(item) if isinstance(item, Table) else _docx_paragraph_markdown(item)
            for item in document.iter_inner_content()
        ]
        return "\n\n".join(b for b in blocks if b.strip())


_LOADERS: dict[str, AnswerLoader] = {
    ".md": MarkdownAnswerLoader(),
    ".markdown": MarkdownAnswerLoader(),
    ".pdf": PdfAnswerLoader(),
    ".docx": DocxAnswerLoader(),
}


def load_answer(path: Path) -> str:
    """Le um arquivo de resposta de Deep Research e retorna seu texto
    (Strategy pattern por formato — md/pdf/docx). Os loaders de pdf/docx
    importam suas dependencias sob demanda para que o pacote base nao
    exija o extra `[ingestion]` apenas para lidar com respostas .md."""
    loader = _LOADERS.get(path.suffix.lower())
    if loader is None:
        raise ExtractionError(f"Formato de resposta nao suportado: {path.suffix!r}")
    return loader.load(path)
