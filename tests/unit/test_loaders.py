from __future__ import annotations

import pytest

pytest.importorskip("docx")

from auditframework.extraction.loaders import DocxAnswerLoader


def _add_hyperlink(paragraph, url: str, text: str):
    """Injeta um `<w:hyperlink>` de verdade no paragrafo (nao existe API
    publica no python-docx para isso) — necessario para reproduzir o bug
    real que motivou a reescrita do loader: `paragraph.runs` pula runs
    dentro de hyperlink, `paragraph.iter_inner_content()` nao."""
    import docx
    from docx.opc.constants import RELATIONSHIP_TYPE
    from docx.oxml import OxmlElement
    from docx.oxml.ns import qn

    part = paragraph.part
    r_id = part.relate_to(url, RELATIONSHIP_TYPE.HYPERLINK, is_external=True)

    hyperlink = OxmlElement("w:hyperlink")
    hyperlink.set(qn("r:id"), r_id)
    run = OxmlElement("w:r")
    t = OxmlElement("w:t")
    t.text = text
    run.append(t)
    hyperlink.append(run)
    paragraph._p.append(hyperlink)


class TestDocxAnswerLoader:
    def test_superscript_run_becomes_bracket_marker(self, tmp_path):
        import docx

        document = docx.Document()
        p = document.add_paragraph("bilheteira")
        sup_run = p.add_run("1")
        sup_run.font.superscript = True
        p.add_run(".")
        path = tmp_path / "resposta.docx"
        document.save(path)

        text = DocxAnswerLoader().load(path)
        assert "bilheteira[1]." in text

    def test_multi_digit_superscript_is_preserved(self, tmp_path):
        import docx

        document = docx.Document()
        p = document.add_paragraph("saturacao")
        sup_run = p.add_run("13")
        sup_run.font.superscript = True
        path = tmp_path / "resposta.docx"
        document.save(path)

        text = DocxAnswerLoader().load(path)
        assert "saturacao[13]" in text

    def test_plain_number_without_superscript_is_not_touched(self, tmp_path):
        import docx

        document = docx.Document()
        document.add_paragraph("Lancado em 2008 com 38 filmes.")
        path = tmp_path / "resposta.docx"
        document.save(path)

        text = DocxAnswerLoader().load(path)
        assert "2008" in text and "[2008]" not in text
        assert "38" in text and "[38]" not in text

    def test_heading_style_becomes_markdown_heading(self, tmp_path):
        import docx

        document = docx.Document()
        document.add_paragraph("Referências citadas", style="Heading 4")
        document.add_paragraph("Corpo normal.")
        path = tmp_path / "resposta.docx"
        document.save(path)

        text = DocxAnswerLoader().load(path)
        lines = [line for line in text.splitlines() if line.strip()]
        assert lines[0] == "#### Referências citadas"
        assert lines[1] == "Corpo normal."

    def test_hyperlink_text_is_preserved_not_dropped(self, tmp_path):
        """Regressao: `paragraph.runs` (usado pelo loader antigo) pula
        runs dentro de `<w:hyperlink>`, perdendo a URL inteira quando ela
        e um hyperlink de verdade (nao texto solto) — confirmado contra o
        .docx real do Gemini antes desta correcao."""
        import docx

        document = docx.Document()
        p = document.add_paragraph("Titulo do artigo, ")
        _add_hyperlink(p, "https://example.com/artigo", "https://example.com/artigo")
        path = tmp_path / "resposta.docx"
        document.save(path)

        text = DocxAnswerLoader().load(path)
        assert "https://example.com/artigo" in text

    def test_table_content_is_included_between_surrounding_paragraphs(self, tmp_path):
        """Regressao: `document.paragraphs` nao inclui paragrafos dentro de
        tabelas — uma tabela inteira (com suas citacoes) desaparecia do
        texto extraido, confirmado contra um .docx real de Deep Research
        cuja tabela comparativa de bilheteria sumia por completo."""
        import docx

        document = docx.Document()
        document.add_paragraph("Paragrafo antes da tabela.")
        table = document.add_table(rows=2, cols=2)
        table.rows[0].cells[0].text = "Filme"
        table.rows[0].cells[1].text = "Bilheteira"
        table.rows[1].cells[0].text = "Avengers: Endgame"
        cell = table.rows[1].cells[1]
        cell.paragraphs[0].text = ""
        p = cell.paragraphs[0]
        p.add_run("2,717B")
        sup_run = p.add_run("1")
        sup_run.font.superscript = True
        document.add_paragraph("Paragrafo depois da tabela.")
        path = tmp_path / "resposta.docx"
        document.save(path)

        text = DocxAnswerLoader().load(path)

        before_idx = text.index("Paragrafo antes da tabela.")
        table_idx = text.index("Avengers: Endgame")
        after_idx = text.index("Paragrafo depois da tabela.")
        assert before_idx < table_idx < after_idx
        assert "2,717B[1]" in text

    def test_table_cell_with_multiple_paragraphs_joins_them(self, tmp_path):
        import docx

        document = docx.Document()
        table = document.add_table(rows=1, cols=1)
        cell = table.rows[0].cells[0]
        cell.paragraphs[0].text = "Primeira linha"
        cell.add_paragraph("Segunda linha")
        path = tmp_path / "resposta.docx"
        document.save(path)

        text = DocxAnswerLoader().load(path)
        assert "Primeira linha\n\nSegunda linha" in text

    def test_table_row_with_only_empty_cells_is_skipped(self, tmp_path):
        import docx

        document = docx.Document()
        document.add_paragraph("Antes.")
        table = document.add_table(rows=2, cols=2)
        table.rows[0].cells[0].text = "Dado"
        table.rows[0].cells[1].text = "Valor"
        # linha 1 (indice 1) permanece com celulas vazias de proposito
        document.add_paragraph("Depois.")
        path = tmp_path / "resposta.docx"
        document.save(path)

        text = DocxAnswerLoader().load(path)
        lines = [line for line in text.split("\n\n") if line.strip()]
        assert "" not in lines
        assert "Antes." in text and "Depois." in text

    def test_superscript_inside_hyperlink_is_still_wrapped(self, tmp_path):
        import docx

        document = docx.Document()
        p = document.add_paragraph("texto")
        _add_hyperlink(p, "https://example.com/nota", "1")
        # marca o run recem-criado dentro do hyperlink como superscript
        hyperlink_run = p._p.findall(".//{http://schemas.openxmlformats.org/wordprocessingml/2006/main}r")[-1]
        rpr = hyperlink_run.makeelement(
            "{http://schemas.openxmlformats.org/wordprocessingml/2006/main}rPr", {}
        )
        vert_align = hyperlink_run.makeelement(
            "{http://schemas.openxmlformats.org/wordprocessingml/2006/main}vertAlign",
            {"{http://schemas.openxmlformats.org/wordprocessingml/2006/main}val": "superscript"},
        )
        rpr.append(vert_align)
        hyperlink_run.insert(0, rpr)
        path = tmp_path / "resposta.docx"
        document.save(path)

        text = DocxAnswerLoader().load(path)
        assert "texto[1]" in text
